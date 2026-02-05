# -*- coding: utf-8 -*-
import datetime
import io
import re
from typing import List, Dict, Any

import streamlit as st
from fpdf import FPDF

from docx import Document
from docx.shared import Pt

# =========================================================
# CONFIG
# =========================================================

APP_TITLE = "📝 Zienswijze Generator Vroendaal"
PAGE_TITLE = "Verzet Vroendaal"
PAGE_ICON = "⚖️"

DEADLINE_TEXT = "📅 DEADLINE: 12 FEBRUARI 2026"
DOSSIER_ZAAKNUMMER = "Z2025-00004367"
ZAAKHEADER = f"Zienswijze zaak {DOSSIER_ZAAKNUMMER}"

INTRO_MD = """
Gebruik deze tool om uw zienswijze te genereren.  
De geselecteerde tekstblokken worden **letterlijk** overgenomen in het Word-document.
"""

INSTRUCTIES_KORT = """
**Instructie (kort)**  
1) Vul uw gegevens in.  
2) Kies of dit een **aanvulling op een pro-forma zienswijze** is.  
3) (Optioneel) dien in **onder protest** bij late/ontbrekende stukken.  
4) Selecteer uw bezwaren en download het document.
"""

# =========================================================
# TEKSTBLOKKEN (LETTERLIJK) — uit Teksten generator Zienswijze final.docx
# Structuur: Heading 1 = category, Heading 2 = section, Heading 3 = title, Normal = text
# =========================================================

BLOCKS = [
    {
        "id": 1,
        "category": "Procedurele gebreken",
        "section": "Misleiding Contracten",
        "title": "Fantoomcontract ondermijnt rechtmatigheid besluitvorming",
        "text": """De anterieure overeenkomst is gepresenteerd als gesloten, terwijl deze juridisch nog niet bestond. Op 26 november 2025 meldde het College aan de gemeenteraad dat de overeenkomst was aangegaan, terwijl alleen de ontwikkelaar had getekend en de gemeentelijke handtekening ontbrak. De overeenkomst is pas op 22 december 2025 rechtsgeldig tot stand gekomen. Hierdoor is de raad feitelijk onjuist geïnformeerd en is de actieve inlichtingenplicht geschonden. Desondanks is het ontwerpplan al op 4 december 2025 ter inzage gelegd, in strijd met de Omgevingswet die vereist dat het kostenverhaal vooraf is verzekerd. Gedurende 18 dagen lag het plan zonder juridische en financiële basis ter inzage. Deze misleiding staat niet op zichzelf: essentiële contractbijlagen ontbreken, basisgegevens van de contractpartij zijn onjuist en contractuele afspraken beperken de beleidsvrijheid van de raad. Dit wijst op een onzorgvuldig en onrechtmatig proces.""",
    },
    {
        "id": 2,
        "category": "Procedurele gebreken",
        "section": "Misleiding Contracten",
        "title": "Structurele misleiding van de gemeenteraad",
        "text": """De gemeenteraad is structureel onjuist en onvolledig geïnformeerd over zowel de juridische status, de inhoud als de gevolgen van het plan. Het College heeft op 26 november 2025 gemeld dat de anterieure overeenkomst was aangegaan, terwijl deze pas op 22 december 2025 rechtsgeldig is ondertekend. Hierdoor is de raad misleid over het verzekerd zijn van het kostenverhaal en is het ontwerpplan vanaf 4 december 18 dagen onrechtmatig ter inzage gelegd. Daarnaast zijn essentiële contractbijlagen met financiële risico’s achtergehouden en bevat de overeenkomst basale fouten, zoals een niet-bestaand KvK-nummer. Ook is ten onrechte gecommuniceerd dat het plan is verkleind, terwijl de bouwmassa feitelijk gelijk bleef. Door contractuele koppeling van grondverkoop aan dit bouwvolume heeft de gemeente een financieel belang gekregen, waardoor de beleidsvrijheid van de raad ernstig is beperkt.""",
    },
    {
        "id": 3,
        "category": "Procedurele gebreken",
        "section": "Misleiding Contracten",
        "title": "Slordige contractvorming ondermijnt besluitvorming",
        "text": """Het foutieve KvK-nummer in de anterieure overeenkomst toont de ernstige onzorgvuldigheid waarmee het contract tot stand is gekomen. In het contract is een niet-bestaand KvK-nummer opgenomen, terwijl de gemeente heeft nagelaten de identiteit van de contractpartij correct te verifiëren. Dit is geen onschuldige typefout, maar symptomatisch voor een geforceerd proces. Deze slordigheid staat niet op zichzelf: de overeenkomst werd aan de raad gepresenteerd als gesloten, terwijl zij juridisch pas op 22 december 2025 tot stand kwam. Daardoor lag het plan 18 dagen onrechtmatig ter inzage zonder verzekerd kostenverhaal. Bovendien ontbreken essentiële contractbijlagen, waaronder kostenramingen, overdrachtstekeningen en het bodemrapport. Hierdoor heeft de raad geen inzicht gehad in financiële en milieutechnische risico’s. Dit geheel wijst op structurele misleiding en een ondeugdelijke juridische basis van het plan.""",
    },
    {
        "id": 4,
        "category": "Procedurele gebreken",
        "section": "Onrechtmatige Inzage",
        "title": "Onrechtmatige terinzagelegging door ontbrekend kostenverhaal",
        "text": """De terinzagelegging van het plan is gestart zonder te voldoen aan een fundamentele wettelijke voorwaarde. Op 4 december 2025 is het ontwerpplan ter inzage gelegd, terwijl de anterieure overeenkomst pas op 22 december 2025 door de gemeente is ondertekend. Daarmee was het kostenverhaal gedurende 18 dagen niet verzekerd, in strijd met artikel 13.6 van de Omgevingswet. De mededeling in het ontwerpplan dat het kostenverhaal was verzekerd, was in deze periode feitelijk onjuist. Daarnaast zijn raad en burgers voorafgaand aan de inzage misleid door de mededeling dat de overeenkomst reeds was aangegaan. Ook was het dossier inhoudelijk onvolledig, doordat essentiële bijlagen zoals kostenramingen, overdrachtstekeningen en het bodemrapport ontbraken. Hierdoor konden burgers hun zienswijze niet baseren op volledige informatie. De inzageprocedure is daarmee niet rechtsgeldig aangevangen.""",
    },
    {
        "id": 5,
        "category": "Procedurele gebreken",
        "section": "Onrechtmatige Inzage",
        "title": "Incompleet dossier schendt rechtsbescherming burgers",
        "text": """Het ter inzage gelegde dossier was fundamenteel incompleet en heeft de rechtsbescherming van burgers en de controlerende taak van de gemeenteraad ernstig geschaad. Essentiële bijlagen bij de anterieure overeenkomst ontbraken, waaronder de kostenraming voor de aanpassing van de Rijksweg, de overdrachtstekening van gemeentegrond en het bodemrapport over ernstige vervuiling. Hierdoor ontbrak inzicht in financiële, juridische en milieutechnische risico’s. Daarnaast zijn cruciale veiligheidsstukken niet ter inzage gelegd, zoals de vereiste provinciale ontheffing in een grondwaterbeschermingsgebied, het veiligheids- en saneringsplan, een bezonningsstudie en relevante ecologische informatie. Ook bestonden er verschillende openbare dossiers en zijn documenten pas na ingebrekestelling laat vrijgegeven. Door deze informatie-asymmetrie is de inzageprocedure niet rechtsgeldig aangevangen en zijn burgers feitelijk belemmerd in het indienen van een volwaardige zienswijze.""",
    },
    {
        "id": 6,
        "category": "Procedurele gebreken",
        "section": "Onrechtmatige Inzage",
        "title": "Schending fair play en obstructie openbaarheid",
        "text": """De gemeente heeft het beginsel van fair play en het recht op een gelijk speelveld ernstig geschonden door doelbewust informatie achter te houden en openbaarmaking te traineren. Interne communicatie laat zien dat ambtenaren zijn geïnstrueerd om geen overleg te voeren met omwonenden om te voorkomen dat zij “gratis munitie” zouden krijgen voor hun zienswijze. Hiermee heeft de gemeente zich niet neutraal opgesteld, maar als tegenpartij van burgers. Daarnaast zijn Woo-verzoeken naar cruciale beslisinformatie bewust vertraagd en pas na ingebrekestelling op 29 januari 2026 mondjesmaat vrijgegeven, slechts twee weken voor het einde van de zienswijzetermijn. Burgers werden hierdoor geconfronteerd met een onwerkbare hoeveelheid complexe stukken. Bovendien is de Woo misbruikt om documenten achter te houden die op grond van de Awb verplicht ter inzage hadden moeten liggen. Hierdoor is sprake van informatie-asymmetrie en een niet-rechtsgeldige inspraakprocedure.""",
    },
    {
        "id": 7,
        "category": "Procedurele gebreken",
        "section": "Integriteit en Vooringenomenheid",
        "title": "Bestuurlijke collusie ondermijnt onafhankelijkheid gemeente",
        "text": """De gemeente heeft haar rol als onafhankelijk en neutraal bevoegd gezag verlaten en zich structureel opgesteld als bondgenoot van de projectontwikkelaar. Interne communicatie toont aan dat de gemeente zichzelf positioneerde als één front met de ontwikkelaar tegenover bewoners en bewust overleg vermeed om hun rechtspositie te verzwakken. Daarnaast is sprake van financiële verstrengeling: de gemeente fungeerde als betaald projectbureau en koppelde grondverkoop contractueel aan dit specifieke bouwvolume, waardoor een direct financieel belang ontstond bij goedkeuring van het plan. Ook zijn ambtelijke processen aangepast aan wensen van de ontwikkelaar, waaronder het cosmetisch herschrijven van negatieve deskundigenadviezen, gezamenlijke persregie en procedurele versnelling op verzoek van de ontwikkelaar. Bewoners werden structureel uitgesloten, terwijl de ontwikkelaar ongehinderde toegang tot bestuur en besluitvorming had. Dit alles vormt een ernstige schending van het verbod op vooringenomenheid en tast de integriteit van de besluitvorming fundamenteel aan.""",
    },
    {
        "id": 8,
        "category": "Procedurele gebreken",
        "section": "Integriteit en Vooringenomenheid",
        "title": "Bewuste boycot van burgerdialoog",
        "text": """De gemeente heeft de dialoog met de Werkgroep doelbewust geboycot als onderdeel van een ambtelijke en bestuurlijke strategie. Interne e-mails tonen aan dat overleg met bewoners expliciet is verboden om te voorkomen dat informatie “tegen ons en de ontwikkelaar” zou worden gebruikt. Wettelijke informatievoorziening werd daarmee gezien als het verschaffen van “gratis munitie” aan een tegenpartij, wat duidt op vooringenomenheid en een schending van artikel 2:4 Awb. De gemeente identificeerde zich openlijk met de ontwikkelaar en verloor haar neutraliteit als bevoegd gezag. Bewoners werden operationeel uitgesloten van gesprekken, informatieavonden en bestuurlijk overleg, terwijl de ontwikkelaar onbeperkte toegang tot het stadhuis behield. Tegelijk werd intern gezocht naar manieren om inhoudelijke discussie in de raad te vermijden. Deze systematische uitsluiting bewijst dat geen sprake was van participatie, maar van een gecoördineerde verdediging van het bouwplan.""",
    },
    {
        "id": 9,
        "category": "Procedurele gebreken",
        "section": "Integriteit en Vooringenomenheid",
        "title": "Betaalplanologie ondermijnt bestuurlijke onafhankelijkheid",
        "text": """De besluitvorming is gestuurd door perverse financiële prikkels die de onafhankelijkheid van de gemeente ernstig aantasten. Via de plankostenscan brengt de gemeente €152.702 in rekening bij de ontwikkelaar voor circa 1.000 ambtelijke uren, waardoor zij feitelijk als betaald projectbureau opereert en een direct financieel belang krijgt bij het doorgaan van het plan. Daarnaast is de verkoop van 266 m² gemeentelijk openbaar groen contractueel gekoppeld aan dit specifieke bouwvolume. Hierdoor ontstaat een situatie van financiële gijzeling: aanpassing van het plan leidt tot contractbreuk of verlies van inkomsten. De democratische afweging wordt zo vooraf financieel vastgezet. Door deze constructies functioneert de gemeente niet langer als onpartijdige belangenafweger, maar als financiële partner van de ontwikkelaar. Dit verklaart de vooringenomen houding tegenover bewoners en vormt een ernstige schending van bestuurlijke integriteit.""",
    },

    {
        "id": 10,
        "category": "Fysieke Leefomgeving",
        "section": "Maat en Schaal",
        "title": "Overschrijding bouwhoogte schaadt maat en schaal",
        "text": """Het bouwplan overschrijdt bewust de toegestane bouwhoogte en past niet binnen de maat en schaal van de wijk Vroendaal. Voor dit gebied geldt volgens de Omgevingsvisie een maximale bouwhoogte van 10 meter, terwijl het ontwerp een totale hoogte van 11,63 meter bereikt, een overschrijding van meer dan 16%. Deze afwijking is gemaskeerd door te rekenen met een kunstmatig peil, waardoor het gebouw op papier lager lijkt dan het in werkelijkheid is. Feitelijk ervaren omwonenden een gevelhoogte van ruim 10,8 meter exclusief opbouw. Daarnaast is de massaliteit verhuld door een onjuiste footprint te presenteren: de werkelijke verstening is circa 50% groter dan aangegeven. In een laagbouwwijk leidt dit tot een ernstige schaalbreuk, verlies van winterzon en forse privacy-aantasting door directe inkijk. Dit plan negeert de menselijke maat en berust op rekenkundige trucs.""",
    },
    {
        "id": 11,
        "category": "Fysieke Leefomgeving",
        "section": "Maat en Schaal",
        "title": "Footprint-misleiding verhult buitensporig bouwvolume",
        "text": """Het plan misleidt over de werkelijke footprint en maskeert daarmee een buitensporig bouwvolume. In de toelichting wordt een footprint van 1.475 m² gepresenteerd, terwijl de feitelijke verstening 2.230 m² bedraagt. Dit verschil van ruim 50% ontstaat doordat de halfverdiepte parkeerkelder niet wordt meegerekend. Hierdoor wordt de fysieke impact van het gebouw structureel onderschat. Een footprint van deze omvang veroorzaakt een ernstige schaalbreuk in de fijnmazige stadsrandwijk Vroendaal en staat haaks op de menselijke maat. Daarnaast leidt de maximale verstening tot verhoogde hittestress, wateroverlast en het onmogelijk maken van wadi’s en een parkachtige inrichting. Het overmaatse volume past bovendien niet op het eigen perceel en wordt alleen mogelijk gemaakt door de verkoop van 266 m² openbaar groen. Deze cijfermatige manipulatie verhult een plan dat de omgeving blijvend aantast.""",
    },
    {
        "id": 12,
        "category": "Fysieke Leefomgeving",
        "section": "Maat en Schaal",
        "title": "Schaalbreuk met laagbouwwijk Vroendaal",
        "text": """Het bouwplan veroorzaakt een fundamentele schaalbreuk met de bestaande laagbouwwijk Vroendaal. De wijk bestaat overwegend uit grondgebonden woningen met een fijnmazige, dorpse structuur, terwijl het plan voorziet in twee massieve, gestapelde bouwblokken van drie lagen op een halfverdiepte parkeerkelder. Deze typologie en massa sluiten niet aan op het bestaande stedenbouwkundige weefsel en vormen een vreemd lichaam in de wijk. Daarnaast wordt het beleidsmatige uitgangspunt voor stadsrandwijken van maximaal 10 meter bouwhoogte overschreden met een feitelijke hoogte van 11,63 meter. Ook de werkelijke footprint van 2.230 m² is aanzienlijk groter dan gepresenteerd en disproportioneel voor de locatie. De visuele dominantie leidt tot ernstige privacy-aantasting door directe inkijk vanuit zestien appartementen op vier bestaande woningen. Van zorgvuldige inpassing is geen sprake.""",
    },
    {
        "id": 13,
        "category": "Fysieke Leefomgeving",
        "section": "Milieu en gezondheid",
        "title": "Onverantwoorde woningbouw op vervuilde locatie",
        "text": """De bouwlocatie betreft een zwaar vervuilde voormalige autosloperij en vormt een acuut risico voor volksgezondheid en milieu. Er is sprake van extreme loodverontreiniging ver boven de interventiewaarde, mogelijke asbest- en PAK-vervuiling en onvolledig bodemonderzoek, terwijl de bestemming wijzigt naar wonen. Bodemrapporten zijn ondanks interne bezwaren toch gebruikt, wat wijst op manipulatie om het plan door te zetten. Tijdens de bouw bestaat een onopgelost veiligheidsconflict tussen asbest- en explosievenprotocollen, zonder integraal veiligheidsplan voor omwonenden op korte afstand. De locatie ligt bovendien in een grondwaterbeschermingsgebied; diepe heipalen kunnen leiden tot verspreiding van vervuiling richting drinkwater, terwijl een vereiste provinciale ontheffing ontbreekt. Sanering is financieel en procedureel niet geborgd: er is geen goedgekeurd saneringsplan en geen budget gereserveerd. Woningbouw onder deze omstandigheden is onverantwoord en strijdig met de zorgplicht.""",
    },
    {
        "id": 14,
        "category": "Fysieke Leefomgeving",
        "section": "Milieu en gezondheid",
        "title": "Genegeerde gezondheidsrisico’s door locatiekeuze",
        "text": """De locatiekeuze brengt ernstige en onvoldoende onderzochte gezondheidsrisico’s met zich mee. Het plan ligt direct aan de Rijksweg, een drukke verkeersader die fungeert als aanzienlijke bron van fijnstof. De gevolgen voor luchtkwaliteit en gezondheid zijn onvoldoende meegewogen, terwijl het project zelf extra verkeer genereert. Daarnaast wordt het risico van landbouwgif genegeerd: op circa 200 meter afstand vinden bespuitingen plaats, zonder dat onderzoek is gedaan naar spuitzones en blootstelling. Deze risico’s stapelen zich op bij de bestaande bodemvervuiling van de voormalige autosloperij. Desondanks voorziet het plan in sociale huurwoningen, bedoeld voor kwetsbare groepen, die juist extra bescherming verdienen. Zonder harde garanties over schone lucht en een veilige leefomgeving worden toekomstige bewoners blootgesteld aan een cumulatie van milieubelasting. Vanuit het oogpunt van volksgezondheid is deze locatiekeuze onverantwoord.""",
    },
    {
        "id": 15,
        "category": "Fysieke Leefomgeving",
        "section": "Milieu en gezondheid",
        "title": "Hittestress door verstening bedreigt leefbaarheid",
        "text": """Het bouwplan verergert hittestress en creëert een onleefbaar microklimaat, in strijd met de zorgplicht voor een gezonde leefomgeving. In de planstukken wordt erkend dat de gevoelstemperatuur ter plaatse richting 2050 kan oplopen tot 47°C. Het ontwerp versterkt dit risico door maximale verstening: twee massieve gebouwen met halfverdiepte parkeerkelders resulteren in een werkelijke footprint van circa 2.230 m² en veroorzaken een sterk hitte-eilandeffect. Natuurlijke bodemkoelte gaat verloren en de voorgestelde wadi’s zijn onvoldoende om dit te compenseren. Daarnaast wordt bestaand volwassen groen gekapt, waardoor schaduw en verkoeling verdwijnen. Herplant biedt pas na decennia effect. De gemeente accepteert deze ingreep ten gunste van bouwlogistiek en financieel gemak. Daarmee wordt een erkend gezondheidsrisico vergroot in plaats van beperkt, wat strijdig is met beleid en zorgplicht.""",
    },
    {
        "id": 16,
        "category": "Fysieke Leefomgeving",
        "section": "Milieu en gezondheid",
        "title": "Gemanipuleerde geluidsrapporten ondermijnen woonklimaat",
        "text": """De geluidsrapportages vormen geen betrouwbare basis voor besluitvorming en brengen de gezondheid van bewoners in gevaar. Interne gemeentelijke adviseurs hebben de onderzoeken herhaaldelijk als “niet akkoord” en oncontroleerbaar bestempeld, omdat essentiële invoergegevens ontbraken of afgekeurde rapporten opnieuw zijn ingediend. Desondanks zijn deze stukken gebruikt in het ontwerpbesluit. Daarnaast zijn rekenkundige trucs toegepast om de geluidsbelasting kunstmatig te verlagen, zoals het meenemen van blinde gevels, het weglaten van de A2 uit de berekeningen en het toepassen van een onjuiste meetmethode voor nieuwbouw. Ook is geen rekening gehouden met extra rolgeluid van grasbetontegels, waardoor geluidshinder structureel wordt onderschat. Toetsing aan geldende milieuzonering ontbrak. Hierdoor ontstaat een schijnveiligheid die het daadwerkelijke woon- en leefklimaat niet beschermt.""",
    },
    {
        "id": 17,
        "category": "Fysieke Leefomgeving",
        "section": "Waterhuishouding",
        "title": "Grondwaterbescherming maakt plan onuitvoerbaar",
        "text": """De bouwlocatie ligt in het Grondwaterbeschermingsgebied Mergelland en vormt een onaanvaardbaar risico voor de drinkwatervoorziening. Het plan voorziet in diepe heipalen en een parkeerkelder die de beschermende kleilaag doorboren. Hierdoor ontstaat hydrologische kortsluiting, waarbij ernstige bodemvervuiling kan doorsijpelen naar het diepe grondwaterreservoir. Dit risico wordt gebagatelliseerd, terwijl vergelijkingen met omliggende bebouwing feitelijk onjuist zijn. Daarnaast is het plan in strijd met de Provinciale Omgevingsverordening, die grondroering dieper dan 3 meter verbiedt zonder ontheffing. Deze ontheffing ontbreekt, waardoor het plan juridisch onuitvoerbaar is. Ook de waterhuishouding is onoplosbaar: infiltratie is verboden, waterberging schiet tekort en het risico op wateroverlast wordt afgewenteld op omwonenden. De veiligheid van het grondwater is niet geborgd.""",
    },
    {
        "id": 18,
        "category": "Fysieke Leefomgeving",
        "section": "Waterhuishouding",
        "title": "Onopgelost waterbergingstekort ondermijnt plan",
        "text": """Het plan kent een structureel waterbergingstekort van 185 m³ dat niet is opgelost. Interne stukken tonen aan dat dit tekort door de gemeente is aangemerkt als “huiswerk voor de ontwikkelaar”, terwijl een omgevingsplan niet kan worden vastgesteld zonder een sluitende wateroplossing. Infiltratie is juridisch verboden vanwege de ligging in het grondwaterbeschermingsgebied en technisch onmogelijk door de slechte doorlatendheid van de bodem. Desondanks wordt in de planstukken ten onrechte uitgegaan van voldoende infiltratiecapaciteit. De voorgestelde wadi’s zijn fysiek niet realiseerbaar door ruimtegebrek naast de kelderwanden, waardoor een badkuip-effect en wateroverlast ontstaan. Het Waterschap waarschuwt voor waterstanden tot 100 cm. De ontwikkelaar voorkomt dit voor zichzelf door ophoging, waarmee de overlast wordt afgewenteld op omwonenden. De waterhuishouding is daarmee technisch en juridisch ondeugdelijk.""",
    },
    {
        "id": 19,
        "category": "Fysieke Leefomgeving",
        "section": "Waterhuishouding",
        "title": "Rioolverlegging maakt plan juridisch onuitvoerbaar",
        "text": """De noodzakelijke rioolverlegging vormt een onoplosbare technische en juridische klem. Door de omvang van het gebouw en de parkeerkelder moet het riool worden verlegd naar circa 4 meter diepte, terwijl in het grondwaterbeschermingsgebied grondroering dieper dan 3 meter strikt is verboden zonder provinciale ontheffing. Deze ontheffing ontbreekt, waardoor het plan juridisch niet uitvoerbaar is. Daarnaast is de voorgestelde oplossing civieltechnisch afgekeurd door gemeentelijke specialisten vanwege knikken in een hoofdleiding en ernstige beheer- en onderhoudsrisico’s. Ondanks deze afwijzingen schuift de gemeente de uitvoerbaarheid door naar de toekomst door te suggereren dat proefsleuven later uitkomst moeten bieden. Daarmee wordt een omgevingsplan vastgesteld zonder zekerheid over de technische haalbaarheid. Dit bevestigt dat het bouwvolume te groot is voor de locatie en dat vitale infrastructuur en milieuregels worden opgeofferd aan commerciële belangen.""",
    },

    {
        "id": 20,
        "category": "Mobiliteit en parkeren",
        "section": "Verkeersveiligheid",
        "title": "Onveilige wijkontsluiting door willekeurige verkeerskeuze",
        "text": """De ontsluiting van het plan via smalle woonstraten is onveilig, onnodig en fysiek onuitvoerbaar. Het grootste deel van het verkeer wordt door een rustige 30-km-woonwijk geleid, terwijl volledige ontsluiting via de Rijksweg niet is onderzocht. Hiermee wordt verkeer van een gebiedsontsluitingsweg verplaatst naar erftoegangswegen, in strijd met verkeersveiligheidsprincipes. De rijbaanbreedte van circa 3,6 meter voldoet niet aan de eis van 4,5 meter vrije doorgang voor hulpdiensten, waardoor bereikbaarheid bij calamiteiten niet is gegarandeerd. Blokkades door pakketbezorgers of foutparkeren maken de situatie extra risicovol. Interne waarschuwingen van deskundigen zijn genegeerd en vervangen door cosmetische teksten om het besluit te rechtvaardigen. Een technisch haalbaar alternatief dat de wijk ontlast, is niet serieus onderzocht. Hiermee wordt een aantoonbaar onveilige ontsluiting afgedwongen.""",
    },
    {
        "id": 21,
        "category": "Mobiliteit en parkeren",
        "section": "Verkeersveiligheid",
        "title": "Manipulatie verkeersrapporten schaadt veiligheid",
        "text": """De verkeersrapportages zijn systematisch gemanipuleerd om een politiek gewenste ontsluiting te legitimeren, ten koste van de verkeersveiligheid. Interne e-mails tonen aan dat deskundigen zijn gevraagd om “mooie zinnen” te formuleren om een technisch onwenselijke route door de woonwijk goed te praten. Fysieke onmogelijkheden zijn genegeerd, zoals een rijbaanbreedte van 3,6 meter die niet voldoet aan de vereiste 4,5 meter voor hulpdiensten. Daarnaast zijn negatieve interne adviezen, waaronder expliciete “niet akkoord”-oordelen over de inrit en verkeerscapaciteit, buiten het besluit gehouden. Ook is technische onkunde gemaskeerd door relevante verkeersbronnen buiten beschouwing te laten. Veiliger alternatieven zijn niet serieus onderzocht of actief gefrustreerd. Hierdoor berust de besluitvorming op een papieren werkelijkheid die de veiligheid van bewoners en hulpdiensten niet waarborgt.""",
    },
    {
        "id": 22,
        "category": "Mobiliteit en parkeren",
        "section": "Verkeersveiligheid",
        "title": "Hulpdiensten fysiek onbereikbaar",
        "text": """De bereikbaarheid voor hulpdiensten is fysiek en wiskundig onmogelijk en vormt een ernstig veiligheidsrisico. De brandweer vereist een vrije doorrijbreedte van 4,5 meter, terwijl de betrokken woonstraat slechts 3,6 meter breed is. Dit tekort van 0,9 meter maakt het onmogelijk om aan de wettelijke veiligheidseisen te voldoen. De situatie wordt verder verslechterd door pakketbezorgers en foutparkeren, waardoor de weg regelmatig volledig wordt geblokkeerd. De bereikbaarheid van hulpdiensten wordt daarmee afhankelijk van toeval. Interne waarschuwingen dat ontsluiting via deze woonstraat onwenselijk is, zijn genegeerd en vervangen door cosmetische teksten om de situatie veilig te laten lijken. Een veiliger alternatief via de hoofdweg is niet onderzocht. Hiermee wordt een aantoonbaar onveilige ontsluiting afgedwongen, waarbij de veiligheid van bewoners ondergeschikt is gemaakt aan andere belangen.""",
    },
    {
        "id": 23,
        "category": "Mobiliteit en parkeren",
        "section": "Parkeerproblematiek",
        "title": "Structureel parkeertekort veroorzaakt wijkoverlast",
        "text": """Het plan voldoet niet aan de geldende parkeernormen en leidt tot directe afwenteling van parkeerdruk op de omgeving. Voor Gebouw A zijn minimaal 44 bewonersplaatsen vereist, terwijl er na toewijzing aan bezoekers slechts 28 beschikbaar blijven. Dit resulteert in een structureel tekort van ten minste 15 bewonersplaatsen. Daarnaast zijn diverse parkeerplaatsen te smal en voldoen zij niet aan de CROW-richtlijnen, waardoor zij feitelijk onbruikbaar zijn voor bezoekers. Ook worden bezoekersplaatsen achter slagbomen meegeteld, terwijl deze juridisch niet openbaar toegankelijk zijn. Hierdoor ontstaat een papieren parkeerbalans die in de praktijk niet functioneert. De voorgenomen verkoop van bezoekersplaatsen ondermijnt het principe van gedeeld gebruik verder. Bezoekers en bewoners zullen daardoor uitwijken naar omliggende woonstraten, in strijd met het uitgangspunt dat nieuwbouw haar parkeerbehoefte volledig op eigen terrein moet oplossen.""",
    },
    {
        "id": 24,
        "category": "Mobiliteit en parkeren",
        "section": "Parkeerproblematiek",
        "title": "Fictieve bezoekersplaatsen achter slagboom",
        "text": """De parkeerbalans is ondeugdelijk doordat bezoekersplaatsen achter slagbomen worden meegeteld. Deze plaatsen zijn juridisch niet openbaar toegankelijk en mogen volgens vaste jurisprudentie niet gelden als bezoekersparkeerplaatsen. Door de fysieke afsluiting functioneren zij feitelijk als privaat terrein voor bewoners. Daarnaast vormt een slagboom een praktische en psychologische drempel: bezoekers en bezorgdiensten zullen niet wachten of aanbellen, maar uitwijken naar de openbare straat. Hierdoor blijft de theoretische capaciteit in de garages onbenut en neemt de parkeerdruk in de omliggende woonstraten toe. Dit leidt tot afwenteling van parkeerproblemen op de wijk, in strijd met het uitgangspunt dat nieuwbouw zijn parkeerbehoefte op eigen terrein moet oplossen. De voorgenomen verkoop van bezoekersplaatsen versterkt dit probleem en ondermijnt de collectieve norm. Zonder vrij toegankelijke bezoekersplaatsen is de parkeerbalans feitelijk onhoudbaar.""",
    },
    {
        "id": 25,
        "category": "Mobiliteit en parkeren",
        "section": "Parkeerproblematiek",
        "title": "Te smalle parkeervakken veroorzaken papiertekort",
        "text": """De ingetekende parkeervakken voldoen niet aan de geldende CROW-richtlijnen en zijn daardoor functioneel onbruikbaar. Bezoekersplaatsen zijn uitgevoerd met een breedte van 2,50 meter, terwijl voor wisselende gebruikers een minimale breedte van 2,65 tot 2,70 meter vereist is. Deze normschending leidt tot zogenoemd schuwgedrag: bezoekers mijden de te krappe vakken of parkeren over de belijning, waardoor aangrenzende plaatsen eveneens onbruikbaar worden. Hierdoor ontstaat een papieren parkeerbalans die in de praktijk niet functioneert. Omdat deze vakken technisch niet geschikt zijn voor bezoekers, mogen zij niet worden meegeteld, wat resulteert in een feitelijk parkeertekort. Bezoekers zullen uitwijken naar de openbare ruimte, waardoor parkeerdruk wordt afgewenteld op omliggende woonstraten. Dit is in strijd met het uitgangspunt dat nieuwbouw haar parkeerbehoefte volledig en deugdelijk op eigen terrein moet oplossen.""",
    },

    {
        "id": 26,
        "category": "Participatie",
        "section": "Schijnparticipatie",
        "title": "Misleidend participatieverslag ondergraaft inspraak",
        "text": """Het participatieverslag van 16 juni 2025 geeft een onjuist en gekleurd beeld van het participatieproces en misleidt de gemeenteraad. Het verslag schetst een constructieve en positieve sfeer, terwijl de bijeenkomsten feitelijk werden gekenmerkt door brede weerstand, frustratie en een unaniem afwijzen van het bouwvolume, ondersteund door 421 handtekeningen. Kritische vragen, toezeggingen en inhoudelijke bezwaren zijn selectief weggelaten of herschreven tot algemene opmerkingen, waarna standaardantwoorden zijn toegevoegd die niet zijn gegeven. Participatie had bovendien geen invloed op kernkeuzes zoals bouwhoogte en massa, die al vastlagen. Alternatieven zijn niet serieus onderzocht en bewoners zijn onder druk gezet. De gemeente heeft dit door de ontwikkelaar opgestelde verslag zonder verificatie overgenomen en daarmee haar vergewisplicht verzaakt. Hierdoor is sprake van schijnparticipatie en besluitvorming op basis van een onjuiste voorstelling van zaken.""",
    },
    {
        "id": 27,
        "category": "Participatie",
        "section": "Schijnparticipatie",
        "title": "Kernbezwaren structureel genegeerd",
        "text": """De participatieprocedure heeft niet geleid tot daadwerkelijke invloed van bewoners en is verworden tot schijnparticipatie. Gedurende het hele traject zijn drie kernbezwaren consequent genegeerd: het te grote bouwvolume, de excessieve bouwhoogte en de onveilige verkeersafwikkeling. De gepresenteerde “verkleining” van het plan betreft slechts een administratieve wijziging; de fysieke massa en impact zijn ongewijzigd gebleven. Ook de bouwhoogte is niet substantieel aangepast en verkeersbezwaren zijn afgedaan met beloften van nader onderzoek. Inbreng van bewoners werd structureel afgewezen zodra deze de businesscase van de ontwikkelaar raakte. Het Burgeralternatief is zonder inhoudelijke toets verworpen en zelfs ter beoordeling voorgelegd aan de ontwikkelaar zelf. De gemeente heeft weerstand gebagatelliseerd en de dialoog actief geblokkeerd. Daarmee is sprake van informeren en negeren, niet van echte participatie.""",
    },
]

# =========================================================
# HELPERS
# =========================================================

def normalize_ws(s: str) -> str:
    return re.sub(r"\s+", " ", s or "").strip()

def build_indexes(blocks: List[Dict[str, Any]]):
    by_id = {b["id"]: b for b in blocks}
    categories = sorted({b["category"] for b in blocks}, key=lambda x: x.lower())
    cat_map = {}
    for cat in categories:
        cat_map[cat] = {}
        sections = sorted({b["section"] for b in blocks if b["category"] == cat}, key=lambda x: x.lower())
        for sec in sections:
            ids = [b["id"] for b in blocks if b["category"] == cat and b["section"] == sec]
            cat_map[cat][sec] = ids
    return by_id, categories, cat_map

def matches_search(block: Dict[str, Any], q: str) -> bool:
    q = (q or "").strip().lower()
    if not q:
        return True
    return (q in block["title"].lower()) or (q in block["text"].lower()) or (q in block["section"].lower())

def create_pdf_bytes(text: str) -> bytes:
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=12)
    pdf.set_font("Arial", size=11)
    clean_text = (text or "").encode("latin-1", "replace").decode("latin-1")
    for line in clean_text.split("\n"):
        pdf.multi_cell(0, 6, line)
    return pdf.output(dest="S").encode("latin-1")

BY_ID, CATEGORIES, CAT_MAP = build_indexes(BLOCKS)

# =========================================================
# DOCX GENERATOR (deterministisch: tekstblokken letterlijk)
# =========================================================

def create_zienswijze_doc(
    data: Dict[str, str],
    selected_points: List[Dict[str, str]],
    is_pro_forma: bool,
    is_protest: bool,
    eigen_bezwaren: List[str],
) -> Document:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    # 1. Header (afzender)
    p_sender = doc.add_paragraph()
    p_sender.add_run(f"{data['naam']}\n").bold = True
    p_sender.add_run(f"{data['adres']}\n{data['postcode']} {data['woonplaats']}\n{data['email']}")

    doc.add_paragraph("\n")

    # Ontvanger + zaakheader
    p_receiver = doc.add_paragraph()
    p_receiver.add_run("Aan de Gemeenteraad van Maastricht\nPostbus 1992\n6201 BZ Maastricht\n").bold = True

    p_case = doc.add_paragraph()
    p_case.add_run(ZAAKHEADER).bold = True

    datum_str = datetime.date.today().strftime("%d-%m-%Y")
    doc.add_paragraph(f"\nMaastricht, {datum_str}")

    # 2. Betreft
    p_subject = doc.add_paragraph()
    p_subject.add_run("Betreft: ").bold = True
    if is_pro_forma:
        subject_text = f"AANVULLING OP PRO-FORMA ZIENSWIJZE ({DOSSIER_ZAAKNUMMER})"
    else:
        subject_text = f"ZIENSWIJZE TAM-OMGEVINGSPLAN 'VROENDAAL' ({DOSSIER_ZAAKNUMMER})"
    p_subject.add_run(subject_text)

    # 3. Inleiding (dynamisch)
    doc.add_paragraph("Geachte leden van de Raad,\n")

    intro_text = ""
    if is_pro_forma:
        intro_text += "Hierbij dien ik, ondergetekende, mijn inhoudelijke aanvulling in op de eerder door mij ingediende pro-forma zienswijze. "
    else:
        intro_text += "Hierbij dien ik, ondergetekende, mijn zienswijze in tegen het ontwerp TAM-Omgevingsplan Vroendaal. "

    if is_protest:
        intro_text += (
            "Ik dien deze zienswijze in ONDER PROTEST.\n\n"
            "Omdat het dossier incompleet is en (cruciale) stukken laat of gefaseerd beschikbaar zijn gesteld, "
            "maak ik een formeel voorbehoud om deze zienswijze later nog aan te vullen zodra alle stukken volledig ter inzage zijn gelegd.\n"
        )

    intro_text += "\nIk breng de volgende bezwaren naar voren:"
    doc.add_paragraph(intro_text)

    # 4. Punten (doorlopende nummering) — letterlijk
    n = 0
    for b in selected_points:
        n += 1
        p = doc.add_paragraph()
        runner = p.add_run(f"\n{n}. {b['title']}")
        runner.bold = True
        doc.add_paragraph(b["text"])

    # Eigen bezwaren (optioneel; doorlopende nummering)
    eigen_clean = [normalize_ws(x) for x in (eigen_bezwaren or []) if normalize_ws(x)]
    if eigen_clean:
        doc.add_paragraph("\nEigen bezwaren (door indiener)\n")
        for item in eigen_clean:
            n += 1
            p = doc.add_paragraph()
            runner = p.add_run(f"\n{n}. {item}")
            runner.bold = True

    # 5. Slot
    slot_text = (
        "\nConclusie & Vordering\n"
        "Ik verzoek u deze zienswijze te betrekken bij de besluitvorming en het plan niet in deze vorm vast te stellen. "
        "Ik verzoek u tevens de ontvangst van deze zienswijze schriftelijk te bevestigen.\n"
    )
    doc.add_paragraph(slot_text)

    doc.add_paragraph("\n\nHoogachtend,\n\n")
    doc.add_paragraph("(handtekening)\n")
    doc.add_paragraph(f"{data['naam']}")
    doc.add_paragraph(f"{data['adres']}\n{data['postcode']} {data['woonplaats']}")

    return doc

# =========================================================
# STREAMLIT UI
# =========================================================

st.set_page_config(page_title=PAGE_TITLE, page_icon=PAGE_ICON, layout="wide")
st.title(APP_TITLE)
st.markdown(INTRO_MD)
st.warning(DEADLINE_TEXT)
st.markdown(INSTRUCTIES_KORT)

# Sidebar: zoek & reset
with st.sidebar:
    st.header("Selectie & hulp")
    search_q = st.text_input("Zoek in titels/tekst", value="", placeholder="bijv. waterberging, parkeertekort, Woo…")
    st.divider()
    if st.button("Selectie wissen", use_container_width=True):
        for b in BLOCKS:
            st.session_state[f"cb_{b['id']}"] = False
        st.rerun()

# STAP 1: Gegevens
st.header("1. Uw Gegevens")
col1, col2 = st.columns(2)
with col1:
    naam = st.text_input("Naam & Voorletters")
    adres = st.text_input("Straat & Huisnummer")
with col2:
    postcode = st.text_input("Postcode")
    woonplaats = st.text_input("Woonplaats", value="Maastricht")
email = st.text_input("E-mailadres")

# STAP 2: Type indiening
st.header("2. Type Indiening")
is_pro_forma = st.checkbox("Dit is een AANVULLING op een eerdere pro-forma zienswijze", value=False)
is_protest = st.checkbox("Ik dien dit in ONDER PROTEST (vanwege late/ontbrekende stukken)", value=True)

# STAP 3: Bezwaren kiezen (mindmap-structuur)
st.header("3. Kies uw Bezwaren")
st.info("Vink aan wat op u van toepassing is. Gebruik de zoekfunctie links om sneller te filteren.")

tabs = st.tabs(CATEGORIES)

for tab, cat in zip(tabs, CATEGORIES):
    with tab:
        st.markdown(f"### {cat}")
        sections = list(CAT_MAP[cat].keys())

        # secties standaard OPEN (zodat niemand denkt dat er “maar 3” zijn)
        for sec in sections:
            ids = CAT_MAP[cat][sec]
            ids_filtered = [i for i in ids if matches_search(BY_ID[i], search_q)]
            if not ids_filtered:
                continue

            with st.expander(f"{sec} ({len(ids_filtered)} punten)", expanded=True):
                for bid in ids_filtered:
                    b = BY_ID[bid]
                    checked = st.checkbox(b["title"], key=f"cb_{bid}", value=st.session_state.get(f"cb_{bid}", False))
if checked:
    st.markdown(b["text"])
    st.markdown("---")

# Teller geselecteerde bezwaren
selected_ids = [b["id"] for b in BLOCKS if st.session_state.get(f"cb_{b['id']}", False)]
st.info(f"🔢 Aantal geselecteerde bezwaren: **{len(selected_ids)}**")

# Eigen bezwaren
st.subheader("Eigen bezwaren / argumenten (optioneel)")
st.caption("U kunt tot 5 eigen punten toevoegen.")

if "custom_items" not in st.session_state:
    st.session_state.custom_items = [""]

cbtn1, cbtn2, cbtn3 = st.columns([1, 1, 2])
with cbtn1:
    if st.button("➕ Voeg punt toe"):
        if len(st.session_state.custom_items) < 5:
            st.session_state.custom_items.append("")
        st.rerun()
with cbtn2:
    if st.button("➖ Verwijder laatste"):
        if len(st.session_state.custom_items) > 1:
            st.session_state.custom_items.pop()
        st.rerun()
with cbtn3:
    st.caption(f"Aantal eigen punten: {len(st.session_state.custom_items)}/5")

custom_inputs = []
for idx in range(len(st.session_state.custom_items)):
    val = st.text_area(
        f"Eigen bezwaar {idx+1}",
        value=st.session_state.custom_items[idx],
        height=80,
        key=f"custom_{idx}",
        placeholder="Schrijf hier uw eigen bezwaar (liefst concreet en feitelijk)."
    )
    custom_inputs.append(val)
st.session_state.custom_items = custom_inputs

# STAP 4: Genereren & Download + instructies
st.header("4. Download & Indieninstructies")

selected_points = [{"title": BY_ID[i]["title"], "text": BY_ID[i]["text"]} for i in selected_ids]

generate_docx = st.button("🚀 Genereer Zienswijze (.docx)", type="primary", use_container_width=True)

if generate_docx:
    if not naam or not adres:
        st.error("Vul naam en adres in.")
        st.stop()

    if len(selected_points) == 0 and not any(normalize_ws(x) for x in st.session_state.custom_items):
        st.error("Kies minimaal één punt (of voeg een eigen bezwaar toe).")
        st.stop()

    user_data = {
        "naam": naam,
        "adres": adres,
        "postcode": postcode,
        "woonplaats": woonplaats,
        "email": email,
    }

    doc = create_zienswijze_doc(
        data=user_data,
        selected_points=selected_points,
        is_pro_forma=is_pro_forma,
        is_protest=is_protest,
        eigen_bezwaren=st.session_state.custom_items,
    )

    bio = io.BytesIO()
    doc.save(bio)

    st.success("Uw document is klaar!")
    safe_name = re.sub(r"[^a-zA-Z0-9_-]+", "_", naam.strip())[:40] or "Zienswijze"
    st.download_button(
        label="⬇️ Download Word-bestand",
        data=bio.getvalue(),
        file_name=f"Zienswijze_{safe_name}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True,
    )

    # Optionele eenvoudige PDF
    try:
        txt = []
        txt.append(f"{naam}\n{adres}\n{postcode} {woonplaats}\n{email}\n")
        txt.append("Aan de Gemeenteraad van Maastricht\nPostbus 1992\n6201 BZ Maastricht\n")
        txt.append(f"{ZAAKHEADER}\n")
        txt.append(f"Maastricht, {datetime.date.today().strftime('%d-%m-%Y')}\n")
        if is_pro_forma:
            txt.append(f"Betreft: AANVULLING OP PRO-FORMA ZIENSWIJZE ({DOSSIER_ZAAKNUMMER})\n")
        else:
            txt.append(f"Betreft: ZIENSWIJZE TAM-OMGEVINGSPLAN 'VROENDAAL' ({DOSSIER_ZAAKNUMMER})\n")
        txt.append("Geachte leden van de Raad,\n")
        for idx, p in enumerate(selected_points, 1):
            txt.append(f"\n{idx}. {p['title']}\n{p['text']}\n")

        pdf_bytes = create_pdf_bytes("\n".join(txt))
        st.download_button(
            label="📄 Download (eenvoudige) PDF",
            data=pdf_bytes,
            file_name=f"Zienswijze_{safe_name}.pdf",
            mime="application/pdf",
            use_container_width=True,
        )
    except Exception:
        pass

st.markdown("---")
st.subheader("📮 Hoe dien ik dit in?")
st.markdown(f"""
U heeft 3 opties om uw zienswijze in te dienen. **Doe dit uiterlijk 12 februari 2026.**

**Optie 1: Digitaal (Snelst)**
- Ga naar de website van de gemeente Maastricht en zoek op 'Zienswijze indienen'.
- Log in met uw DigiD.
- Upload het Word-bestand (of sla het eerst op als PDF).

**Optie 2: Per Post (Aangetekend aanbevolen)**
- Print het document uit.
- **Zet uw handtekening** onderaan de brief.
- Stuur het naar:  
  > Gemeenteraad van Maastricht  
  > Postbus 1992  
  > 6201 BZ Maastricht  

**Optie 3: Per E-mail (alleen als dit passend is in uw situatie)**
- Mail naar `post@maastricht.nl` o.v.v. "Zienswijze zaak {DOSSIER_ZAAKNUMMER}" en vraag om een ontvangstbevestiging.

*Tip: Stuur ook een kopie naar de griffie (`griffie@maastricht.nl`) zodat raadsleden weten dat u gereageerd heeft.*
""")
